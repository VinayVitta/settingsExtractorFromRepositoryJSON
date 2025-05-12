# Tasks Summary
# Total Number of tasks
# Total Running state tasks
# Total LogStream Tasks
# Total Apply Changes Tasks
# Total Store Changes Tasks
# Total Apply and Store Changes tasks
# Total Tasks without LogStream

# Count of child tasks for each LogStream

# Batch Tuning with MIN/MAX/Memory settings
# For new adds -
# Import the query
# Add execution part like for batch_tuning_queries
# Add printing part


import pandas as pd
import duckdb
from helpers.queries import tasksCounts, changeProcessTuning, handlingPolicy, logStream
from tabulate import tabulate
from docx import Document
from docx.shared import Inches


def read_csv(path):
    data_df = pd.read_csv(path)
    # data_columns = data_df.columns
    return data_df


def summarize_tasks(query_list, table_name="data_df", show=True):
    """Execute multiple DuckDB SQL summary queries and return combined result as a DataFrame."""
    results = []
    for q in query_list:
        results.append(duckdb.query(q).to_df())
    final_df = pd.concat(results, ignore_index=True)
    if show:
        print(final_df.to_string(index=False))  # No ugly boxes
    return final_df


def export_tables_to_word(summary_tables, filename="task_summary.docx", title="Task Summary Report"):
    doc = Document()
    doc.add_heading(title, 0)

    for table_info in summary_tables:
        doc.add_heading(table_info["title"], level=1)

        if "notes" in table_info:
            doc.add_paragraph(f"Notes: {table_info['notes']}")

        df = table_info["df"]
        if df.empty:
            doc.add_paragraph("No data available.")
            continue

        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid'

        # Header
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)

        # Rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, col in enumerate(df.columns):
                row_cells[i].text = str(row[col])

        doc.add_paragraph()  # spacing

    doc.save(filename)


if __name__ == "__main__":
    csv_file_path = r"C:\Users\VIT\OneDrive - QlikTech Inc\QlikVit\Customers\EdwardJones\PlatformReview\filecloud-20250430192131\run_output_20250507_094558\exportRepositoryCSV_20250507_094558.csv"
    data_df = read_csv(csv_file_path)
    # Task Counts
    task_counts_queries = [
        tasksCounts.total_tasks_query,
        tasksCounts.running_tasks_query,
        tasksCounts.replication_tasks_query,
        tasksCounts.logstream_tasks_query,
        tasksCounts.apply_changes_query,
        tasksCounts.store_changes_query,
        tasksCounts.apply_store_changes_query,
        tasksCounts.no_logstream_query
    ]
    # Batch Tuning
    batch_tuning_queries = [changeProcessTuning.batch_tuning]

    # Transaction Offloading
    transactionOffloading_queries = [changeProcessTuning.transaction_memory]

    # Control Tables
    control_tables_queries = [changeProcessTuning.control_tables]

    # LOB Size
    lob_size_query = [changeProcessTuning.lob_size]

    # Handling Policy
    dll_handlingPolicy_query = [handlingPolicy.ddl_handling]

    # Overall Policy
    handling_policy_query  = [handlingPolicy.handling_policy]

    # Error Handling =
    error_handling_query = [handlingPolicy.error_handling]

    #LogStream
    totalChildTasksForEachParent_query = [logStream.totalChildTasksForEachParent]

    # Running Queries
    task_counts_summary_df = pd.concat([duckdb.query(q).to_df() for q in task_counts_queries], ignore_index=True)
    batch_tuning_summary_df = pd.concat([duckdb.query(q).to_df() for q in batch_tuning_queries], ignore_index=True)
    transactionOffloading_summary_df = pd.concat([duckdb.query(q).to_df() for q in transactionOffloading_queries], ignore_index=True)
    control_tables_df = pd.concat([duckdb.query(q).to_df() for q in control_tables_queries], ignore_index=True)
    lob_size_df = pd.concat([duckdb.query(q).to_df() for q in lob_size_query], ignore_index=True)
    ddl_handling_df = pd.concat([duckdb.query(q).to_df() for q in dll_handlingPolicy_query], ignore_index=True)
    handling_policy_df = pd.concat([duckdb.query(q).to_df() for q in handling_policy_query], ignore_index=True)
    error_handling_df = pd.concat([duckdb.query(q).to_df() for q in error_handling_query], ignore_index=True)
    totalChildTasksForEachParent_df = pd.concat([duckdb.query(q).to_df() for q in totalChildTasksForEachParent_query], ignore_index=True)

    # Print Result
    print(task_counts_summary_df.to_string(index=False))
    print(batch_tuning_summary_df.to_string(index=False))
    print(transactionOffloading_summary_df.to_string(index=False))
    print(control_tables_df.to_string(index=False))
    print(lob_size_df.to_string(index=False))
    print(ddl_handling_df.to_string(index=False))
    print(handling_policy_df.to_string(index=False))
    print(error_handling_df.to_string(index=False))
    print(totalChildTasksForEachParent_df.to_string(index=False))
    # print(tabulate(summary_df, headers='keys', tablefmt='grid'))
    summary_tables = []

    summary_tables.append({
        "title": "Running Tasks Summary",
        "notes": "Shows distinct running tasks and total tables in running state.",
        "df": task_counts_summary_df
    })

    # print(summary_tables)
    # export_tables_to_word(summary_tables, "full_task_summary.docx")




